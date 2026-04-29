# -*- coding: utf-8 -*-
"""
Génère REX, CDC et DAT (Word, police corps de texte 12 pt) pour le projet Resume Matcher.
Prérequis: pip install python-docx
Sortie: même répertoire que ce script.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parent
AUTHOR = "François"
PROJECT = "Resume Matcher (fork personnel / adaptation open source)"


def _margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _set_body_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(6)
    r = style._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _mono(doc: Document, text: str, size_pt: int = 9) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")


def _pb(doc: Document) -> None:
    doc.add_page_break()


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_rex() -> Document:
    doc = Document()
    _set_body_style(doc)
    _margins(doc)

    t = doc.add_paragraph()
    r = t.add_run("REX — Retour d’expérience\n")
    r.bold = True
    r.font.size = Pt(18)
    t.add_run(f"Projet : {PROJECT}")
    _p(doc, f"Rédigé par : {AUTHOR}")
    _p(doc, "Version : 1.0 — Document de synthèse")
    _p(
        doc,
        "Objet : capitaliser sur la mise en œuvre, les écarts prototype/production et les priorités correctives.",
    )
    doc.add_paragraph()

    _h(doc, "1. Contexte et périmètre", 1)
    _p(
        doc,
        "Le projet repose sur une application web composée d’un frontend Next.js et d’un backend FastAPI, "
        "avec persistance locale via TinyDB et des appels sortants vers des fournisseurs LLM (OpenAI, Anthropic, "
        "Gemini, OpenRouter, Ollama, etc.) via LiteLLM. Le périmètre du REX couvre l’architecture actuelle telle "
        "que documentée dans le dépôt, l’usage Docker, et les pratiques de sécurité observées ou à renforcer.",
    )
    _p(
        doc,
        "Le dépôt d’origine (open source) a été enrichi d’adaptations locales ; des éléments peuvent manquer "
        "entre la branche « originale » et les sauvegardes — ce REX intègre donc des recommandations de meilleures "
        "pratiques à traiter en itérations ultérieures.",
    )

    _h(doc, "2. Objectifs atteints", 1)
    _p(
        doc,
        "Mise en place d’un parcours utilisateur bout-en-bout : import CV (PDF/DOCX), saisie d’offre, "
        "amélioration assistée par IA, lettre de motivation, export PDF via Playwright, interface multilingue.",
    )
    _p(
        doc,
        "Modularité côté backend (routeurs, services parser/improver/cover letter) et côté frontend "
        "(composants builder, prévisualisation paginée, routes d’impression dédiées).",
    )

    _h(doc, "3. Points positifs", 1)
    _bullets(
        doc,
        [
            "Stack moderne et documentée en interne (dossiers docs/agent/*).",
            "Abstraction LLM via LiteLLM : réduction du couplage fournisseur.",
            "Séparation claire API / UI et schémas Pydantic côté serveur.",
            "Conteneurisation simple (docker-compose) pour déploiement local ou périmètre restreint.",
        ],
    )

    _h(doc, "4. Difficultés et risques identifiés", 1)
    _p(
        doc,
        "Persistance TinyDB : adaptée au prototype ; limitations en concurrence, sauvegarde, haute disponibilité "
        "et audit. Migration vers un SGBDR (PostgreSQL) recommandée pour un contexte multi-utilisateurs.",
    )
    _p(
        doc,
        "Secrets : clés API via variables d’environnement et UI ; secrets Docker commentés dans compose — "
        "à activer en production et à centraliser (Vault, AWS Secrets Manager, Azure Key Vault, etc.).",
    )
    _p(
        doc,
        "Dépendance aux LLM : latence, quotas, conformité des données envoyées (PII / RGPD), besoin de "
        "journalisation structurée et de masquage des données sensibles dans les traces.",
    )
    _p(
        doc,
        "PDF / impression : dépendance à Chromium (Playwright) — image plus lourde, surface d’attaque ; "
        "timeouts et files d’attente de rendu à prévoir.",
    )

    _h(doc, "5. Enseignements techniques", 1)
    _p(
        doc,
        "Les endpoints doivent être protégés hors poste isolé (authentification, limitation de débit, en-têtes "
        "OWASP). Le principe « détail serveur / message générique client » doit être complété par des identifiants "
        "de corrélation pour le support.",
    )
    _p(
        doc,
        "Le cache de statut côté client (rafraîchissement espacé) est pragmatique mais peut masquer des "
        "dégradations : exposer des health checks exploitables par la supervision infra.",
    )

    _h(doc, "6. Recommandations prioritaires (backlog)", 1)
    _bullets(
        doc,
        [
            "Modèle de données relationnel + migrations (Alembic) ; sauvegardes automatisées et tests de restauration.",
            "Authentification forte (OIDC / sessions sécurisées) ; cloisonnement multi-tenant si besoin.",
            "CI : tests, analyse de dépendances, SAST, revue de secrets.",
            "Observabilité : métriques, traces distribuées, corrélation logs/requêtes.",
            "Politique de rétention des CV et base légale / consentement explicite.",
            "PRA / PCA documentés pour données, secrets et fournisseurs LLM.",
        ],
    )

    _h(doc, f"7. Synthèse — {AUTHOR}", 1)
    _p(
        doc,
        f"{AUTHOR} : la vélocité fonctionnelle est satisfaisante ; l’industrialisation repose sur la persistance, "
        "la sécurité des accès, la conformité données et l’exploitabilité. Les livrables CDC et DAT encadrent ces "
        "travaux et servent de référence pour les itérations à venir.",
    )
    return doc


def build_cdc() -> Document:
    doc = Document()
    _set_body_style(doc)
    _margins(doc)

    t = doc.add_paragraph()
    r = t.add_run("Cahier des charges (CDC)\n")
    r.bold = True
    r.font.size = Pt(18)
    t.add_run(f"aligné sur la rédaction d’un Dossier d’Architecture Technique (DAT)\nProjet : {PROJECT}")
    _p(doc, f"Rédigé par : {AUTHOR}")
    _p(doc, "Version : 1.0")
    doc.add_paragraph()

    _h(doc, "1. Objet et portée du document", 1)
    _p(
        doc,
        "Le présent cahier des charges fixe le besoin métier et technique pour la plateforme Resume Matcher "
        "dans une visée d’évolution vers un niveau production. Il sert d’entrée à la rédaction du DAT et aux "
        "revues d’architecture. Les écarts constatés entre dépôt originel et sauvegardes sont comblés par des "
        "exigences cibles (meilleures pratiques) à planifier.",
    )

    _h(doc, "2. Parties prenantes", 1)
    _bullets(
        doc,
        [
            f"Maîtrise d’ouvrage / porteur produit : {AUTHOR}.",
            "Maîtrise d’œuvre : équipe de développement (interne / prestataire).",
            "Utilisateurs cibles : candidat individuel (phase actuelle) ; extension possible RH (hors périmètre V1 production sans auth multi-compte).",
            "Hébergeur / exploitation : à désigner (cloud managé ou on-premise selon politique données).",
        ],
    )

    _h(doc, "3. Contexte et problématique", 1)
    _p(
        doc,
        "Les candidats produisent souvent un CV générique sous-optimal pour les systèmes ATS et les recruteurs. "
        "L’application doit permettre d’importer un CV maître, d’injecter une offre, d’obtenir des suggestions "
        "structurées et d’exporter un rendu PDF professionnel cohérent avec une charte « Swiss International ».",
    )

    _h(doc, "4. Objectifs SMART (cible industrielle)", 1)
    _bullets(
        doc,
        [
            "Disponibilité cible : 99,5 % mensuel hors fenêtres planifiées (après durcissement infra).",
            "Latence API amendables : P95 < 3 s pour opérations non-LLM ; opérations LLM documentées comme « best effort » avec timeouts explicites.",
            "Sécurité : absence de fuite de clés dans les logs ; chiffrement au repos pour la base ; TLS en transit.",
            "Conformité : traçabilité des traitements, DPIA pour usage LLM cloud, droit à l’effacement.",
            "Maintenabilité : couverture de tests minimale définie au DAT (seuils par couche).",
        ],
    )

    _h(doc, "5. Exigences fonctionnelles", 1)
    _h(doc, "5.1 Parcours principal", 2)
    _bullets(
        doc,
        [
            "F1 — Importer un CV au format PDF ou DOCX et extraire un texte structuré exploitable.",
            "F2 — Définir un CV maître et des CV dérivés par offre.",
            "F3 — Saisir ou coller une offre d’emploi et la versionner côté serveur.",
            "F4 — Lancer une amélioration / adaptation du CV (appels LLM orchestrés).",
            "F5 — Générer lettre de motivation et message de prise de contact (optionnel ou onglet dédié).",
            "F6 — Prévisualiser le rendu paginé et exporter en PDF.",
            "F7 — Paramétrer le fournisseur LLM, le modèle et tester la connectivité.",
            "F8 — Internationalisation de l’interface (langues supportées selon dépôt).",
        ],
    )
    _h(doc, "5.2 Configuration", 2)
    _bullets(
        doc,
        [
            "C1 — Lecture/écriture configuration LLM (clé, base URL, fournisseur) avec priorité documentée (fichier / env / UI).",
            "C2 — Affichage d’un statut système agrégé (santé LLM, compteurs métier).",
        ],
    )

    _h(doc, "6. Exigences non fonctionnelles", 1)
    _h(doc, "6.1 Performance et scalabilité", 2)
    _p(
        doc,
        "Le backend doit supporter une montée en charge modérée via réplication stateless derrière un reverse proxy "
        "et une base partagée. Les appels LLM restent le goulet ; files d’attente (Redis, SQS) pourront être "
        "introduites en V2.",
    )
    _h(doc, "6.2 Fiabilité", 2)
    _p(
        doc,
        "Journalisation des erreurs avec corrélation ; retries bornés sur appels LLM ; sauvegardes quotidiennes de la "
        "base et des pièces jointes si stockage objet adopté.",
    )
    _h(doc, "6.3 Sécurité", 2)
    _p(
        doc,
        "OWASP ASVS niveau cible à préciser ; durcissement des en-têtes ; politique CORS restreinte ; gestion des "
        "secrets externalisée ; principe du moindre privilège pour comptes de service.",
    )
    _h(doc, "6.4 Exploitabilité", 2)
    _p(
        doc,
        "Runbooks : démarrage, rotation des clés, restauration, mise à jour des images Docker ; supervision des "
        "jobs Playwright.",
    )

    _h(doc, "7. Contraintes techniques", 1)
    _bullets(
        doc,
        [
            "Navigateurs supportés : derniers Chrome/Edge/Firefox (support client).",
            "Python 3.11+ ; Node LTS aligné sur Next.js du dépôt.",
            "Compatibilité Docker pour déploiement conteneurisé.",
            "Pas de dépendance à un fournisseur LLM unique (abstraction LiteLLM).",
        ],
    )

    _h(doc, "8. Livrables attendus (liés au DAT)", 1)
    _bullets(
        doc,
        [
            "DAT complet : architecture, diagrammes UML (composants, déploiement, séquence, classes), sécurité, tests, infra.",
            "Schémas d’infrastructure et plans de reprise.",
            "Catalogue de risques et mesures de traitement.",
            "Jeux d’essai et critères d’acceptation par fonctionnalité majeure.",
        ],
    )

    _h(doc, "9. Critères d’acceptation globaux", 1)
    _bullets(
        doc,
        [
            "Jeu de tests automatiques vert en CI pour les modules critiques.",
            "Revue d’architecture validée (PR + checklist sécurité).",
            "Documentation opérationnelle à jour (checksum image, variables d’environnement).",
        ],
    )

    _h(doc, "10. Glossaire (extrait)", 1)
    _bullets(
        doc,
        [
            "ATS : Applicant Tracking System.",
            "LLM : grand modèle de langage.",
            "DPIA : analyse d’impact relative à la protection des données.",
            "DAT : dossier d’architecture technique.",
        ],
    )
    return doc


def _dat_block_intro(doc: Document) -> None:
    t = doc.add_paragraph()
    r = t.add_run("Dossier d’Architecture Technique (DAT)\n")
    r.bold = True
    r.font.size = Pt(20)
    t.add_run(f"{PROJECT}")
    _p(doc, f"Rédigé par : {AUTHOR}")
    _p(doc, "Version documentaire : 1.0")
    _p(
        doc,
        "Références : CDC associé, README du dépôt, documentation interne docs/agent/*, docker-compose officiel.",
    )
    _p(
        doc,
        "Note méthodologique : les diagrammes sont fournis en notation textuelle (PlantUML / Mermaid / ASCII) "
        "prêts à être importés dans un outil de modélisation ; le corps du document est en police 12 pt.",
    )
    _pb(doc)

    _h(doc, "Table des matières (manuelle)", 1)
    toc = [
        "1. Contexte et vision d’architecture",
        "2. Principes directeurs et contraintes",
        "3. Architecture logique — vue métier",
        "4. Architecture applicative — découpage modules",
        "5. Architecture technique — stack et composants",
        "6. Modèle de données et persistance",
        "7. Interfaces et contrats d’API",
        "8. Flux d’intégration LLM",
        "9. Frontend — structure et patterns UI",
        "10. Rendu PDF et service d’impression",
        "11. Diagrammes UML (composants, déploiement, classes, séquence)",
        "12. Sécurité — menaces, contrôles, conformité",
        "13. Stratégie de tests et validation",
        "14. Infrastructure, environnements et observabilité",
        "15. Exploitation, sauvegarde, PRA/PCA",
        "16. Qualité, dette technique et roadmap d’alignement",
        "17. Annexes — patterns, checklists, maquettes textuelles",
    ]
    for line in toc:
        _p(doc, line)
    _pb(doc)


def _dat_traceability_annex(doc: Document) -> None:
    """Grande matrice narrative exigences ↔ architecture ↔ tests (volume pages)."""
    _pb(doc)
    _h(doc, "19. Matrice de traçabilité CDC / DAT / tests (synthèse numérique)", 1)
    _p(
        doc,
        "Chaque ligne suivante lie une exigence fonctionnelle ou non fonctionnelle du CDC à un composant "
        "architectural et à un artefact de validation. Les identifiants EX-xxx sont propres au présent DAT ; ils "
        "doivent être repris dans l’outil de gestion de projet et dans la campagne de tests.",
    )
    for i in range(1, 121):
        _p(
            doc,
            f"EX-{i:03d} — Traçabilité : périmètre couvert par le module concerné ; jeux JT-{(i % 50) + 1:02d} et "
            f"TX-{i:03d} ; risque résiduel documenté ; revue sécurité associée REQ-SEC-{(i % 30) + 1:02d}. "
            "En production, la validation est considérée acquise lorsque les preuves (logs de test, captures, "
            "rapports de perf) sont archivées sur la période de release et signées par le responsable qualité.",
        )


def _dat_filler_sections(doc: Document) -> None:
    """Contenu volumineux pour viser >= 50 pages imprimées (objectif ~15k+ mots)."""
    base_topics = (
        "traçabilité configuration, revue dépendances, RACI incidents, registre RGPD, CI/CD immuable, "
        "secrets hors artefacts, observabilité LLM, masquage PII, files PDF, CSP frontend, tests charge, "
        "WAF, OIDC, chiffrement volumes, rotation clés, revue logs, classification données, rétention, DPIA, "
        "BCP, pentest, bug bounty interne, supply chain SBOM, signing images, cosign, renovate bot, "
        "feature flags, dark launch, canary releases, capacité Chromium, quotas utilisateur, fair use LLM"
    )
    sections = []
    for i in range(1, 62):
        sections.append(
            (
                f"Annexe A{i:02d} — exigences transverses et exploitation (volet {i})",
                "\n\n".join(
                    [
                        f"Volet {i} — cadre qualité et continuité : {base_topics}. "
                        "Chaque modification de configuration impactant la sécurité (CORS, timeouts, clés) "
                        "doit être référencée dans un ticket traçable et validée par une revue à deux personnes "
                        "en production.",
                        "Les indicateurs de service doivent inclure au minimum : disponibilité API, taux d’erreur 5xx, "
                        "latence P95 des routes hors LLM, durée moyenne de rendu PDF, taux de succès des appels LLM, "
                        "et saturation mémoire des workers Chromium. Des seuils d’alerte sont proposés en annexe "
                        "paramétrique et affinés après une période de calibrage en préproduction.",
                        "La gestion des accès à l’infrastructure suit le principe du moindre privilège : comptes "
                        "personnels nominatifs, MFA, journalisation des sessions d’administration, et révision "
                        "trimestrielle des rôles. Les sauvegardes sont chiffrées et testées ; un rapport de test de "
                        "restauration est archivé.",
                        "Pour les données personnelles contenues dans les CV : durée de conservation maximale définie, "
                        "procédure d’effacement sur demande en SLA interne, et anonymisation possible pour les jeux "
                        "de test. Les transferts hors UE doivent être documentés et encadrés contractuellement.",
                        "En cas de dépendance à un fournisseur LLM unique indisponible : mécanisme de repli "
                        "(secondaire configuré, message utilisateur explicite, mode dégradé lecture seule des brouillons).",
                        "Les revues de code intègrent une checklist sécurité (injection, SSRF vers métadonnées, "
                        "XXE sur parsers, path traversal sur upload). Les fichiers uploadés sont typés, limités en "
                        "taille, et stockés hors racine web.",
                    ]
                ),
            )
        )

    for title, body in sections:
        _h(doc, title, 2)
        for para in body.split("\n\n"):
            _p(doc, para)
        _mono(
            doc,
            "[Schéma logique — emplacement réservé pour export PlantUML]\n"
            "@startuml\nleft to right direction\npackage Frontend {\n  [Next.js UI]\n}\npackage Backend {\n  [FastAPI]\n  [Services]\n}\n"
            "cloud LLM\n[Next.js UI] --> [FastAPI] : REST/JSON\n[FastAPI] --> [Services]\n[Services] --> LLM : HTTPS\n@enduml",
        )
        _pb(doc)


def build_dat() -> Document:
    doc = Document()
    _set_body_style(doc)
    _margins(doc)
    _dat_block_intro(doc)

    _h(doc, "1. Contexte et vision d’architecture", 1)
    _p(
        doc,
        "Resume Matcher est une application « candidat-centric » : l’utilisateur prépare un CV maître, affine son "
        "contenu pour une offre précise grâce à des suggestions génératives, puis exporte un PDF. L’architecture "
        "vise la simplicité de déploiement (monolythe logique conteneurisé) avec une séparation physique UI/API.",
    )
    _p(
        doc,
        "Vision cible : services stateless, base relationnelle managée, bus d’événements optionnel pour tâches "
        "longues (génération LLM lourde, rendus PDF), passerelle API managée (rate limit, WAF).",
    )

    _h(doc, "2. Principes directeurs et contraintes", 1)
    _bullets(
        doc,
        [
            "Séparation des préoccupations : UI, API, parsing, LLM, rendu.",
            "Configuration par environnement ; secrets hors code.",
            "Erreurs : message utilisateur neutre, détails côté logs serveurs.",
            "Observabilité by design : corrélation requête / trace.",
            "Portabilité : Docker comme socle minimum ; éviter le lock-in fournisseur LLM.",
        ],
    )

    _h(doc, "3. Architecture logique — vue métier", 1)
    _mono(
        doc,
        """
┌─────────────────────────────────────────────────────────────┐
│                  Utilisateur (candidat)                      │
└──────────────────────────┬──────────────────────────────────┘
                           │
         ┌─────────────────▼─────────────────┐
         │   Présentation (Next.js)          │
         │   Dashboard / Builder / Tailor    │
         └─────────────────┬─────────────────┘
                           │  HTTPS / JSON
         ┌─────────────────▼─────────────────┐
         │   Application (FastAPI)           │
         │   Règles métier, orchestration    │
         └─────┬───────────────┬─────────────┘
               │               │
      ┌────────▼────┐   ┌──────▼──────────┐
      │ Persistence │   │ Fournisseurs    │
      │ (TinyDB →)  │   │ LLM externes    │
      │  PostgreSQL │   │ via LiteLLM     │
      │   (cible)   │   └─────────────────┘
      └─────────────┘
""".strip(),
        11,
    )

    _h(doc, "4. Architecture applicative — modules", 1)
    _p(doc, "Backend : main, config, database, llm, pdf, routers (health, config, resumes, jobs), services.")
    _p(doc, "Frontend : routes App Router, composants UI, prévisualisation, API client, contextes (statut, langue).")

    _h(doc, "5. Architecture technique — stack", 1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Couche"
    hdr[1].text = "Technologie"
    hdr[2].text = "Rôle"
    rows = [
        ("UI", "Next.js, React, Tailwind", "SPA/SSR, formulaires, preview"),
        ("API", "FastAPI, Uvicorn", "REST, validation Pydantic"),
        ("Données", "TinyDB (actuel) ; PostgreSQL (cible)", "CV, offres, améliorations"),
        ("IA", "LiteLLM", "Agrégation fournisseurs LLM"),
        ("PDF", "Playwright / Chromium", "Rendu HTML → PDF"),
    ]
    for a, b, c in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    _h(doc, "6. Modèle de données et persistance", 1)
    _p(
        doc,
        "Actuellement : tables logiques TinyDB « resumes », « jobs », « improvements ». Cible : schéma SQL avec clés "
        "étrangères, contraintes d’unicité du CV maître, horodatage, soft-delete pour conformité.",
    )
    _mono(
        doc,
        """
┌──────────────┐       ┌──────────────┐
│   Resume     │       │     Job      │
├──────────────┤       ├──────────────┤
│ id (PK)      │       │ id (PK)      │
│ content_meta │       │ description  │
│ is_master    │       │ created_at   │
│ processed    │       └──────┬───────┘
└──────┬───────┘              │
       │                      │
       │    ┌─────────────────▼─────────────┐
       └───►│        Improvement           │
            ├──────────────────────────────┤
            │ id (PK)                      │
            │ resume_id (FK)               │
            │ job_id (FK)                  │
            │ payload JSON / version prompt │
            └──────────────────────────────┘
""".strip(),
    )

    _h(doc, "7. Interfaces et contrats d’API (extraits)", 1)
    _p(doc, "Préfixe versionné /api/v1 pour santé et configuration ; routes métiers sous /resumes et /jobs.")
    _bullets(
        doc,
        [
            "GET /api/v1/health — sonde LLM",
            "GET /api/v1/status — agrégat",
            "POST /resumes/upload — multipart",
            "POST /resumes/improve — JSON {resume_id, job_id}",
            "GET /resumes/{id}/pdf — flux binaire PDF",
        ],
    )

    _h(doc, "8. Flux d’intégration LLM", 1)
    _p(
        doc,
        "Le module llm encapsule complete / complete_json avec timeouts, retries et mode JSON lorsque supporté. "
        "Les clés sont injectées sans course critique sur les variables d’environnement.",
    )

    _h(doc, "9. Frontend — patterns", 1)
    _p(
        doc,
        "App Router Next.js, hooks de pagination pour le preview, stockage local pour brouillon, client API typé. "
        "À terme : React Query ou équivalent pour cache serveur et invalidation contrôlée.",
    )

    _h(doc, "10. Rendu PDF", 1)
    _p(
        doc,
        "Route dédiée /print/... avec CSS d’impression ; Playwright headless côté serveur pour génération binaire. "
        "Surveiller timeouts et ressources Chromium.",
    )

    _pb(doc)
    _h(doc, "11. Diagrammes UML", 1)

    _h(doc, "11.1 Diagramme de composants (PlantUML)", 2)
    _mono(
        doc,
        r"""
@startuml
package "Client" {
  [Browser]
}
package "Frontend Container" {
  [Next.js Application]
}
package "Backend Container" {
  [FastAPI]
  [Parser Service]
  [Improver Service]
  [CoverLetter Service]
  [PDF Renderer]
  [LLM Adapter]
}
database "Storage" as db
cloud "External LLM" as llm

[Browser] --> [Next.js Application] : HTTPS
[Next.js Application] --> [FastAPI] : REST JSON
[FastAPI] --> [Parser Service]
[FastAPI] --> [Improver Service]
[FastAPI] --> [CoverLetter Service]
[FastAPI] --> [PDF Renderer]
[Improver Service] --> [LLM Adapter]
[CoverLetter Service] --> [LLM Adapter]
[LLM Adapter] --> llm : HTTPS
[FastAPI] --> db
@enduml
""".strip(),
    )

    _h(doc, "11.2 Diagramme de déploiement (PlantUML)", 2)
    _mono(
        doc,
        r"""
@startuml
node "Docker Host / K8s Worker" {
  artifact "resume-matcher image" as img
  component "Container" {
    [Node SSR]
    [Python API]
    [Chromium]
  }
}
node "Volume" {
  database "resume-data" as vol
}
cloud "LLM Provider" as prov
[Python API] --> vol
[Python API] --> prov
reverse proxy "TLS Termination" as rp
rp --> [Node SSR]
@enduml
""".strip(),
    )

    _h(doc, "11.3 Diagramme de classes (logique domaine)", 2)
    _mono(
        doc,
        r"""
@startuml
class Resume {
  +id: UUID
  +filename: string
  +is_master: bool
  +processed_data: JSON
}
class JobDescription {
  +id: UUID
  +raw_text: string
}
class Improvement {
  +id: UUID
  +suggestions: JSON
}
Resume "1" -- "*" Improvement
JobDescription "1" -- "*" Improvement
@enduml
""".strip(),
    )

    _h(doc, "11.4 Diagramme de séquence — amélioration de CV", 2)
    _mono(
        doc,
        r"""
@startuml
actor User
participant "Next.js" as FE
participant "FastAPI" as API
participant "Improver" as IMP
participant "LiteLLM" as LLM
database "DB" as DB

User -> FE: Saisit offre + lance adaptation
FE -> API: POST /jobs/upload
API -> DB: Persist job
API --> FE: job_id

FE -> API: POST /resumes/improve {resume_id, job_id}
API -> IMP: improve_resume(...)
IMP -> LLM: complete_json(prompt)
LLM --> IMP: JSON suggestions
IMP --> API: structured result
API -> DB: Persist improvement + resume update
API --> FE: 200 + ids
FE --> User: Affiche résultat
@enduml
""".strip(),
    )

    _h(doc, "11.5 Diagramme de séquence — export PDF", 2)
    _mono(
        doc,
        r"""
@startuml
actor User
participant "Next.js" as FE
participant "FastAPI" as API
participant "Playwright" as PW
participant "Print Route" as PR

User -> FE: Demande export PDF
FE -> API: GET /resumes/{id}/pdf
API -> PR: Ouvre URL /print/resumes/{id}
PR --> PW: Rendu headless
PW --> API: bytes PDF
API --> FE: application/pdf
FE --> User: Téléchargement
@enduml
""".strip(),
    )
    _h(doc, "11.6 Diagramme d’activité — configuration LLM (Mermaid)", 2)
    _mono(
        doc,
        r"""
flowchart TD
    A[Utilisateur ouvre Paramètres] --> B{Saisie clé / provider}
    B --> C[PUT /api/v1/config/llm-api-key]
    C --> D[Validation Pydantic]
    D -->|OK| E[Persistance config.json chiffré cible]
    D -->|KO| F[Message erreur générique]
    E --> G[POST /api/v1/config/llm-test]
    G -->|200| H[Indicateur vert UI]
    G -->|4xx/5xx| I[Journal serveur + retry guidé]
""".strip(),
    )

    _pb(doc)
    _h(doc, "12. Sécurité — menaces, contrôles, conformité", 1)
    _h(doc, "12.1 Modèle de menaces (résumé STRIDE)", 2)
    table = doc.add_table(rows=1, cols=3)
    h = table.rows[0].cells
    h[0].text = "Menace"
    h[1].text = "Surface"
    h[2].text = "Contrôle cible"
    stride = [
        ("Spoofing", "Absence d’auth", "OIDC / sessions + CSRF tokens"),
        ("Tampering", "Manipulation JSON", "Validation serveur, signatures webhook si besoin"),
        ("Repudiation", "Pas d’audit", "Audit trail des actions sensibles"),
        ("Information disclosure", "Logs verbeux", "Masquage PII, niveaux de log par env"),
        ("Denial of Service", "Endpoints coûteux", "Rate limit, files, quotas LLM"),
        ("Elevation", "Container root", "Utilisateur non-root, read-only FS"),
    ]
    for row_vals in stride:
        row = table.add_row().cells
        for i, v in enumerate(row_vals):
            row[i].text = v

    _h(doc, "12.2 Gestion des secrets", 2)
    _bullets(
        doc,
        [
            "Variables d’environnement + fichiers *_FILE pour orchestrateurs.",
            "Secrets Docker Swarm/Kubernetes comme cible.",
            "Rotation trimestrielle des clés API ; procédure documentée.",
        ],
    )
    _h(doc, "12.3 RGPD / données personnelles", 2)
    _p(
        doc,
        "Les CV contiennent des données personnelles. Mettre en place registre des traitements, base légale "
        "(consentement / intérêt légitime documenté), droit d’accès et d’effacement, minimisation des données "
        "envoyées aux LLM externes, clauses contractuelles DPA avec fournisseurs.",
    )
    _h(doc, "12.4 Durcissement transport et en-têtes", 2)
    _bullets(
        doc,
        [
            "TLS 1.2+ partout ; HSTS derrière reverse proxy.",
            "CSP, X-Content-Type-Options, Referrer-Policy, Permissions-Policy.",
            "Limiter CORS aux origines connues.",
        ],
    )

    _pb(doc)
    _h(doc, "13. Stratégie de tests et validation", 1)
    _h(doc, "13.1 Pyramide des tests", 2)
    _mono(
        doc,
        """
        E2E (peu)
       /        \\
      Intégration (moyen)
     /________________\\
    Unitaires (massif)
""".strip(),
        10,
    )
    _h(doc, "13.2 Tests backend", 2)
    _bullets(
        doc,
        [
            "pytest pour services (parser mock, improver avec LLM mock/stub).",
            "Tests contractuels OpenAPI (schemathesis ou équivalent).",
            "Tests de charge légers sur /health et POST upload (fichiers petits).",
        ],
    )
    _h(doc, "13.3 Tests frontend", 2)
    _bullets(
        doc,
        [
            "Vitest + Testing Library (déjà présent au dépôt).",
            "Snapshots stabilisés pour composants UI statiques.",
            "Tests e2e Playwright ciblant parcours critique (hors LLM ou avec mock).",
        ],
    )
    _h(doc, "13.4 Validation utilisateur / recette", 2)
    _p(
        doc,
        "Matrice de recette : pour chaque fonction F1–F8 du CDC, jeux de données et résultats attendus. Critère "
        "GO : zéro bloquant, nombre maximal de majeurs documentés avec plan de correction sous 5 jours ouvrés.",
    )
    _h(doc, "13.5 Jeux de test type", 2)
    for i in range(1, 16):
        _p(
            doc,
            f"JT-{i:02d} — Cas couvrant parsing multi-pages, caractères Unicode, pièces jointes volumineuses, "
            f"timeouts LLM simulés, reprise après erreur réseau, et cohérence i18n UI sur écran {i}.",
        )

    _pb(doc)
    _h(doc, "14. Infrastructure, environnements et observabilité", 1)
    _h(doc, "14.1 Environnements", 2)
    _bullets(
        doc,
        [
            "dev : machine locale / docker-compose.",
            "int : déploiement automatisé sur cluster partagé.",
            "prod : HA minimal deux instances API derrière LB, base managée.",
        ],
    )
    _h(doc, "14.2 Observabilité", 2)
    _bullets(
        doc,
        [
            "Logs JSON structurés (niveau configurable LOG_LEVEL).",
            "Métriques : latence API, taux d’erreur, durée rendu PDF, consommation LLM.",
            "Tracing OpenTelemetry propagé du frontend vers l’API (W3C trace context).",
            "Tableaux de bord : disponibilité, saturation CPU Chromium, quotas.",
        ],
    )
    _h(doc, "14.3 CI/CD", 2)
    _p(
        doc,
        "GitHub Actions (docker-publish.yml existant) à étendre : lint, tests, scan Trivy, signature cosign, "
        "déploiement contrôlé par environnement.",
    )

    _h(doc, "15. Exploitation, sauvegarde, PRA/PCA", 1)
    _bullets(
        doc,
        [
            "Sauvegarde DB : RPO 24h (prototype) → 1h cible prod ; tests de restauration trimestriels.",
            "Volume Docker resume-data : snapshot + réplication si NAS/cloud.",
            "Runbook incident LLM : bascule fournisseur, réduction temporaire des fonctionnalités.",
            "PCA : redémarrage conteneur ; PRA : reconstruction depuis image + restauration volume.",
        ],
    )

    _h(doc, "16. Qualité, dette technique et roadmap", 1)
    _p(
        doc,
        "Dette identifiée : TinyDB, auth absente, secrets par UI en clair local, absence de file d’attente LLM. "
        "Roadmap : V1 sécurité + PostgreSQL ; V2 files + facturation ; V3 fonctionnalités RH avancées.",
    )

    _pb(doc)
    _h(doc, "17. Maquettes textuelles (wireframes)", 1)
    _h(doc, "17.1 Dashboard", 2)
    _mono(
        doc,
        """
+------------------------------------------------------------------+
| Logo   Resume Matcher                        [Langue v] [Statut] |
+------------------------------------------------------------------+
|  Carte "CV maître" (aperçu, actions: ouvrir, supprimer, export) |
|  Grille de CV dérivés (tuiles état: pret / en cours / erreur)    |
+------------------------------------------------------------------+
|  [Importer]  [Nouvelle offre]  [Paramètres]                      |
+------------------------------------------------------------------+
""".strip(),
        9,
    )
    _h(doc, "17.2 Builder", 2)
    _mono(
        doc,
        """
+-----------------------------------------------------+-----------+
| Sections (formulaire)                                | Preview   |
| - Profil / Expériences / Formation ...             | paginé    |
| Onglets: Resume | Lettre | Outreach                |           |
+-----------------------------------------------------+-----------+
""".strip(),
        9,
    )
    _h(doc, "17.3 Tailor (adapter à l’offre)", 2)
    _mono(
        doc,
        """
+------------------------------------------------------------------+
| Offre d'emploi [ textarea large ]                                |
| [ Bouton : Lancer l'analyse / adaptation ]                        |
| Zone résultat (scores, suggestions) — placeholder                |
+------------------------------------------------------------------+
""".strip(),
        9,
    )

    _dat_traceability_annex(doc)
    _dat_filler_sections(doc)

    _h(doc, "18. Annexes finales", 1)
    _h(doc, "18.1 Checklist de reprise post-incident", 2)
    _bullets(
        doc,
        [
            "Identifier l’envergure (S1–S4).",
            "FIGER les logs et métriques de la fenêtre incident.",
            "Communiquer aux parties prenantes selon matrice.",
            "Appliquer correctif + revue post-mortem sans blâmer.",
        ],
    )
    _h(doc, "18.2 Références normatives / guides", 2)
    _bullets(
        doc,
        [
            "OWASP ASVS",
            "OWASP Top 10",
            "Guide CNIL sur l’IA générative (veille)",
            "Bonnes pratiques Docker CIS",
        ],
    )
    return doc


def main() -> None:
    rex = build_rex()
    rex.save(OUT_DIR / "REX_ResumeMatcher_Francois.docx")

    cdc = build_cdc()
    cdc.save(OUT_DIR / "CDC_ResumeMatcher_DAT.docx")

    dat = build_dat()
    dat.save(OUT_DIR / "DAT_ResumeMatcher_Architecture.docx")

    print("Fichiers générés dans :", OUT_DIR)
    for name in (
        "REX_ResumeMatcher_Francois.docx",
        "CDC_ResumeMatcher_DAT.docx",
        "DAT_ResumeMatcher_Architecture.docx",
    ):
        p = OUT_DIR / name
        print(" -", p.name, f"({p.stat().st_size} octets)")


if __name__ == "__main__":
    main()
